import os
import uuid
from fpdf import FPDF
from docx import Document

def generate_file(text, title, file_format):
    output_dir = 'app/static/files'
    os.makedirs(output_dir, exist_ok=True)
    
    # Генерация UUID для имени файла
    unique_filename = f"{uuid.uuid4()}.{file_format}"
    file_path = os.path.join(output_dir, unique_filename)
    
    if file_format == "txt":
        with open(file_path, "w") as file:
            file.write(text)
    elif file_format == "docx":
        document = Document()
        document.add_heading(title, 0)
        document.add_paragraph(text)
        document.save(file_path)
    elif file_format == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt=title, ln=True, align='C')
        pdf.multi_cell(0, 10, text)
        pdf.output(file_path)

    return file_path
